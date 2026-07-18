from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "参赛软件系统文档.md"
OUTPUT = ROOT / "参赛软件系统文档.docx"
ASSETS = ROOT / "assets"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360

BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
LIGHT_BLUE = "EAF2F8"
VERY_LIGHT_BLUE = "F5F9FC"
GRAY = "F2F4F7"
MID_GRAY = "D0D5DD"
TEXT = "24303F"
MUTED = "667085"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(run, latin="Calibri", east_asia="Microsoft YaHei", size=11, bold=None, color=None) -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_widow_control(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:widowControl")) is None:
        p_pr.append(OxmlElement("w:widowControl"))


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "0")
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell)


def widths_for_table(headers: list[str]) -> list[int]:
    n = len(headers)
    joined = " ".join(headers)
    if n == 2:
        return [2500, 6860]
    if n == 3:
        return [1900, 4860, 2600]
    if n == 4:
        return [1450, 3750, 2860, 1300]
    if n == 5:
        if "输入或触发" in joined:
            return [1050, 3450, 1650, 2150, 1060]
        return [1150, 3200, 1900, 2050, 1060]
    return [CONTENT_WIDTH_DXA // n] * n


def add_hyperlink(paragraph, text: str, url: str, color=BLUE) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(fonts)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(r"(`[^`]+`|https?://[^\s]+)")


def add_inline_runs(paragraph, text: str, size=11, color=TEXT, bold=False) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size, color=color, bold=bold)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, latin="Consolas", east_asia="Microsoft YaHei", size=max(8.5, size - 1), color=DEEP_BLUE, bold=False)
        else:
            trailing = ""
            while token and token[-1] in "，。；：,.;:)）":
                trailing = token[-1] + trailing
                token = token[:-1]
            add_hyperlink(paragraph, token, token)
            if trailing:
                run = paragraph.add_run(trailing)
                set_run_font(run, size=size, color=color, bold=bold)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color, bold=bold)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=9, color=MUTED)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "目录将在打开或渲染文档时更新。"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


def add_numbering_definition(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    # Use stable, non-overlapping identifiers.  Recomputing max() on python-docx's
    # cached numbering tree can return the same id for several definitions.
    abstract_id = {"heading": 101, "bullet": 102, "number": 103}[kind]
    num_id = abstract_id

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)
    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        if level > 0:
            restart = OxmlElement("w:lvlRestart")
            restart.set(qn("w:val"), str(level))
            lvl.append(restart)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        if kind == "heading":
            lvl_text.set(qn("w:val"), ["%1", "%1.%2", "%1.%2.%3"][level])
        elif kind == "number":
            lvl_text.set(qn("w:val"), "%{}．".format(level + 1))
        else:
            lvl_text.set(qn("w:val"), ["•", "–", "◦"][level])
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "space")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(360 + level * 360))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        if kind == "heading":
            ind.set(qn("w:left"), "0")
            ind.set(qn("w:hanging"), "0")
        else:
            ind.set(qn("w:left"), str(720 + level * 360))
            ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        lvl.append(p_pr)
        if kind == "bullet":
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), "Arial")
            fonts.set(qn("w:hAnsi"), "Arial")
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.insert(0, num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DEEP_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = RGBColor.from_string(DEEP_BLUE)
    code_style.paragraph_format.left_indent = Inches(0.22)
    code_style.paragraph_format.right_indent = Inches(0.12)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(7)
    code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    if "Figure Caption" not in styles:
        caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Figure Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(9)
    caption.paragraph_format.keep_with_next = True


def set_section_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_document(doc: Document) -> None:
    for section in doc.sections:
        set_section_page(section)
    configure_styles(doc)
    doc.core_properties.title = "基于大模型的个性化资源生成与学习多智能体系统——参赛软件系统文档"
    doc.core_properties.subject = "需求规格、系统设计、测试说明与报告、部署说明、用户手册"
    doc.core_properties.author = "2026 软件杯参赛项目组"
    doc.core_properties.keywords = "软件杯, 个性化学习, 多智能体, SQLite, Python"
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def font_path() -> str:
    candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\msyhbd.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return str(path)
    return "arial.ttf"


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str, width=5) -> None:
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        pts = [(x2, y2), (x2 - direction * 18, y2 - 11), (x2 - direction * 18, y2 + 11)]
    else:
        direction = 1 if y2 > y1 else -1
        pts = [(x2, y2), (x2 - 11, y2 - direction * 18), (x2 + 11, y2 - direction * 18)]
    draw.polygon(pts, fill=color)


def rounded_box(draw, box, fill, outline, title, detail, title_font, detail_font, title_color=DEEP_BLUE):
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = box
    normalized_title_color = title_color if str(title_color).startswith("#") else f"#{title_color}"
    draw.text((x1 + 24, y1 + 17), title, font=title_font, fill=normalized_title_color)
    if detail:
        draw.multiline_text((x1 + 24, y1 + 57), detail, font=detail_font, fill="#475467", spacing=7)


def create_architecture_diagram() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1600, 1080), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    fp = font_path()
    title_font = ImageFont.truetype(fp, 38)
    layer_font = ImageFont.truetype(fp, 28)
    detail_font = ImageFont.truetype(fp, 21)
    small_font = ImageFont.truetype(fp, 18)
    draw.text((80, 50), "个性化学习多智能体系统总体架构", font=title_font, fill="#163B5C")
    draw.text((80, 105), "本地优先 · 关系化存储 · 模型可选 · 判题受限", font=detail_font, fill="#667085")

    layers = [
        ((100, 175, 1500, 300), "桌面启动层", "启动软件.bat / 便携启动器 / 自动端口 / 健康检查 / Edge 应用模式", "#EAF2F8"),
        ((100, 335, 1500, 460), "Web 表现层", "课程创建 · 学习路径 · 掌握度 · 测验复习 · 学习报告 · 学习陪练", "#F5F9FC"),
        ((100, 495, 1500, 620), "应用服务层", "HTTP API · 会话服务 · 方案服务 · 应用状态服务 · 模型适配 · 错误边界", "#EEF4FF"),
        ((100, 655, 1500, 780), "多智能体与领域层", "画像 · 诊断 · 图谱 · 前测 · 路径 · 资源 · 评分 · 治理 · 洞察 · 陪练", "#F0FDF9"),
    ]
    for box, title, detail, fill in layers:
        rounded_box(draw, box, fill, "#7CA9CC", title, detail, layer_font, detail_font)
    for y in (300, 460, 620):
        draw_arrow(draw, (800, y + 4), (800, y + 30), "#7CA9CC", 4)

    rounded_box(draw, (100, 835, 575, 1015), "#F9FAFB", "#98A2B3", "SQLite 数据层", "方案 / 任务 / 测验 / 掌握度\n报告 / 应用状态", layer_font, small_font)
    rounded_box(draw, (625, 835, 1075, 1015), "#FFF7ED", "#F5A45D", "内置 Python 运行器", "CPython 3.13.14\n独立进程 / 资源与能力限制", layer_font, small_font)
    rounded_box(draw, (1125, 835, 1500, 1015), "#F5F3FF", "#9B8AFB", "外部模型（可选）", "OpenAI 兼容 HTTPS\n未配置时本地降级", layer_font, small_font)
    for x in (335, 850, 1312):
        draw_arrow(draw, (x, 780), (x, 827), "#7CA9CC", 4)
    image.save(ASSETS / "architecture.png", quality=95)


def create_learning_loop_diagram() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1600, 940), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    fp = font_path()
    title_font = ImageFont.truetype(fp, 38)
    box_font = ImageFont.truetype(fp, 26)
    detail_font = ImageFont.truetype(fp, 18)
    draw.text((80, 45), "个性化学习闭环", font=title_font, fill="#163B5C")
    draw.text((80, 100), "学习证据持续回流，驱动掌握度、补救资源与后续路径更新", font=detail_font, fill="#667085")

    boxes = [
        ((80, 190, 400, 340), "课程输入", "主题 · 目标 · 基础\n周期 · 偏好", "#EAF2F8", "#7CA9CC"),
        ((475, 190, 795, 340), "画像与知识图谱", "目标分解 · 概念依赖\n学习风险", "#EEF4FF", "#7CA9CC"),
        ((870, 190, 1190, 340), "诊断前测", "覆盖关键概念\n识别薄弱点", "#F5F3FF", "#9B8AFB"),
        ((1265, 190, 1520, 340), "路径与资源", "阶段 · 每日任务\n讲解 · 示例", "#F0FDF9", "#5FAE97"),
        ((1265, 540, 1520, 690), "每日学习", "任务 · 笔记\n项目实践", "#F0FDF9", "#5FAE97"),
        ((870, 540, 1190, 690), "练习与评分", "选择 · 简答 · 代码\n解析 · 错题", "#FFF7ED", "#F5A45D"),
        ((475, 540, 795, 690), "掌握度与补救", "概念证据 · 薄弱项\n定向资源", "#FEF3F2", "#F97066"),
        ((80, 540, 400, 690), "洞察与报告", "进度 · 优势 · 风险\n下一步建议", "#F9FAFB", "#98A2B3"),
    ]
    for box, title, detail, fill, border in boxes:
        rounded_box(draw, box, fill, border, title, detail, box_font, detail_font)
    for a, b in [((400, 265), (467, 265)), ((795, 265), (862, 265)), ((1190, 265), (1257, 265)),
                 ((1392, 340), (1392, 532)), ((1265, 615), (1198, 615)), ((870, 615), (803, 615)),
                 ((475, 615), (408, 615))]:
        draw_arrow(draw, a, b, "#7CA9CC", 5)
    draw.line([(240, 540), (240, 445), (1030, 445), (1030, 350)], fill="#7CA9CC", width=5)
    draw_arrow(draw, (1030, 445), (1030, 350), "#7CA9CC", 5)
    draw.text((430, 405), "学习证据回流，重新诊断并调整后续路径", font=detail_font, fill="#475467")
    draw.rounded_rectangle((80, 770, 1520, 865), radius=15, fill="#F5F9FC", outline="#B7CEE1", width=2)
    draw.text((115, 795), "外部模型可增强内容表达；本地规则引擎保证核心结构与离线演示。所有状态通过方案标识写入 SQLite。", font=detail_font, fill="#344054")
    image.save(ASSETS / "learning-loop.png", quality=95)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    set_section_page(section)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)

    stripe = doc.add_table(rows=1, cols=1)
    stripe.alignment = WD_TABLE_ALIGNMENT.CENTER
    stripe.autofit = False
    set_table_geometry(stripe, [CONTENT_WIDTH_DXA])
    cell = stripe.cell(0, 0)
    set_cell_shading(cell, BLUE)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("2026 软件杯参赛项目  ·  SOFTWARE SYSTEM DOCUMENTATION")
    set_run_font(run, size=9, color=WHITE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(52)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("基于大模型的个性化资源生成与学习多智能体系统")
    set_run_font(run, size=25, color=DEEP_BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(34)
    run = p.add_run("参赛软件系统文档")
    set_run_font(run, size=30, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("需求规格  ·  系统设计  ·  测试说明与报告  ·  部署说明  ·  用户手册")
    set_run_font(run, size=12, color=MUTED, bold=False)

    accent = doc.add_table(rows=1, cols=1)
    accent.alignment = WD_TABLE_ALIGNMENT.LEFT
    accent.autofit = False
    set_table_geometry(accent, [CONTENT_WIDTH_DXA])
    acell = accent.cell(0, 0)
    set_cell_shading(acell, LIGHT_BLUE)
    ap = acell.paragraphs[0]
    ap.paragraph_format.space_before = Pt(7)
    ap.paragraph_format.space_after = Pt(7)
    add_inline_runs(ap, "本地优先 · SQLite 关系化存储 · 内置 CPython · OpenAI 兼容模型增强", size=10.5, color=DEEP_BLUE, bold=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(40)
    meta = doc.add_table(rows=5, cols=2)
    set_table_geometry(meta, [2500, 6860])
    meta.style = "Table Grid"
    values = [
        ("文档编号", "PLS-MAS-DOC-001"),
        ("软件 / 文档版本", "V0.1.0 / V1.0"),
        ("基线分支", "Zip"),
        ("编制日期", "2026-07-17"),
        ("文档状态", "参赛提交版"),
    ]
    for row, (key, value) in zip(meta.rows, values):
        set_cell_shading(row.cells[0], GRAY)
        for idx, text in enumerate((key, value)):
            p = row.cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline_runs(p, text, size=10, color=DEEP_BLUE if idx == 0 else TEXT, bold=idx == 0)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("文档用途")
    set_run_font(run, size=10, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_inline_runs(p, "用于竞赛评审、现场展示、安装部署、验收测试与后续维护。", size=9.5, color=MUTED)


def configure_body_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [7000, 2360])
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(2)
    add_inline_runs(left, "个性化资源生成与学习多智能体系统", size=8.5, color=DEEP_BLUE, bold=True)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(2)
    add_inline_runs(right, "PLS-MAS-DOC-001 · V1.0", size=8.5, color=MUTED)
    p = header.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), MID_GRAY)
    borders.append(bottom)
    p_pr.append(borders)

    footer = section.footer
    p = footer.paragraphs[0]
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:space"), "2")
    top.set(qn("w:color"), MID_GRAY)
    borders.append(top)
    p_pr.append(borders)
    add_page_number(p)

    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), "1")


def add_unnumbered_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=18, color=BLUE, bold=True)


def add_heading(doc: Document, text: str, level: int, page_break_before: bool = False) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.page_break_before = page_break_before
    add_inline_runs(p, text, size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DEEP_BLUE, bold=True)
    set_paragraph_widow_control(p)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.24)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    set_paragraph_widow_control(p)
    add_inline_runs(p, text, size=11, color=TEXT)


def add_list_paragraph(doc: Document, text: str, style: str, num_id: int | None = None, level=0) -> None:
    p = doc.add_paragraph(style=style)
    if num_id is not None:
        apply_numbering(p, num_id, level)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    set_paragraph_widow_control(p)
    add_inline_runs(p, text, size=11, color=TEXT)


def new_numbered_list_id(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    base_num = numbering.num_having_numId(5)
    abstract_id = base_num.abstractNumId.val
    num = numbering.add_num(abstract_id)
    override = num.add_lvlOverride(ilvl=0)
    override.add_startOverride(1)
    return int(num.numId)


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph(style="Code Block")
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), GRAY)
    p_pr.append(shd)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, latin="Consolas", east_asia="Microsoft YaHei", size=8.5, color=DEEP_BLUE)


def parse_table_line(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = [parse_table_line(line) for line in lines]
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    headers = rows[0]
    table = doc.add_table(rows=len(rows), cols=len(headers))
    table.style = "Table Grid"
    widths = widths_for_table(headers)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for r_idx, row_values in enumerate(rows):
        for c_idx, value in enumerate(row_values[: len(headers)]):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, GRAY)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            if r_idx == 0:
                p.paragraph_format.keep_with_next = True
            add_inline_runs(p, value, size=8.5, color=DEEP_BLUE if r_idx == 0 else TEXT, bold=r_idx == 0)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_figure(doc: Document, alt: str, rel_path: str) -> None:
    path = (ROOT / rel_path).resolve()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.28))
    caption = doc.add_paragraph(style="Figure Caption")
    add_inline_runs(caption, alt, size=9, color=MUTED)


def parse_markdown(doc: Document, text: str) -> None:
    lines = text.splitlines()
    i = 0
    paragraph_buffer: list[str] = []
    first_heading = True

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_body_paragraph(doc, " ".join(x.strip() for x in paragraph_buffer))
            paragraph_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            i += 1
            continue
        if stripped == "<!-- PAGEBREAK -->":
            flush_paragraph()
            doc.add_page_break()
            i += 1
            continue
        if stripped.startswith("```"):
            flush_paragraph()
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue
        image_match = re.fullmatch(r"!\[(.+?)\]\((.+?)\)", stripped)
        if image_match:
            flush_paragraph()
            add_figure(doc, image_match.group(1), image_match.group(2))
            i += 1
            continue
        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            add_heading(doc, heading_match.group(2), len(heading_match.group(1)), page_break_before=first_heading)
            first_heading = False
            i += 1
            continue
        if stripped.startswith("|") and "|" in stripped[1:]:
            flush_paragraph()
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_markdown_table(doc, table_lines)
            continue
        bullet_match = re.match(r"^-\s+(.+)$", stripped)
        if bullet_match:
            flush_paragraph()
            while i < len(lines):
                match = re.match(r"^-\s+(.+)$", lines[i].strip())
                if not match:
                    break
                add_list_paragraph(doc, match.group(1), "List Bullet")
                i += 1
            continue
        number_match = re.match(r"^\d+[.)、．]\s*(.+)$", stripped)
        if number_match:
            flush_paragraph()
            list_num_id = new_numbered_list_id(doc)
            while i < len(lines):
                match = re.match(r"^\d+[.)、．]\s*(.+)$", lines[i].strip())
                if not match:
                    break
                add_list_paragraph(doc, match.group(1), "List Number", num_id=list_num_id)
                i += 1
            continue
        paragraph_buffer.append(stripped)
        i += 1
    flush_paragraph()


def build() -> Path:
    create_architecture_diagram()
    create_learning_loop_diagram()
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_page(body_section)
    configure_body_header_footer(body_section)
    add_unnumbered_title(doc, "目录")
    toc = doc.add_paragraph()
    toc.paragraph_format.space_after = Pt(12)
    add_toc(toc)

    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    out = build()
    print(out)
